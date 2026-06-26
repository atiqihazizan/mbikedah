#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_docx_from_markdown(md_file, docx_file):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.style = 'No Spacing'
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines (but add spacing)
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph('_' * 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Handle headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            
            if level == 1:
                p = doc.add_heading(text, level=1)
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.bold = True
            elif level == 2:
                p = doc.add_heading(text, level=2)
                p.runs[0].font.size = Pt(16)
            elif level == 3:
                p = doc.add_heading(text, level=3)
                p.runs[0].font.size = Pt(14)
            elif level == 4:
                p = doc.add_heading(text, level=4)
                p.runs[0].font.size = Pt(12)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.count('|') >= 2:
            table_data = []
            # Collect table rows
            while i < len(lines) and '|' in lines[i]:
                row = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                if row and not all(c in '-|: ' for c in ''.join(row)):  # Skip separator rows
                    table_data.append(row)
                i += 1
            
            if table_data:
                # Create table
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Make header bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            continue
        
        # Handle lists
        if re.match(r'^[-*]\s+', line) or re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^[-*]\s+', '', line)
            text = re.sub(r'^\d+\.\s+', '', text)
            
            # Handle bold in list items
            p = doc.add_paragraph(text, style='List Bullet')
            format_text_with_bold(p, text)
            i += 1
            continue
        
        # Handle checkboxes
        if '✅' in line or '⚠️' in line:
            p = doc.add_paragraph(line, style='List Bullet')
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        format_text_with_bold(p, line)
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Laporan telah ditukar kepada: {docx_file}")

def format_text_with_bold(paragraph, text):
    """Format text with bold markers"""
    # Remove markdown bold markers and apply formatting
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Code/inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Regular text
            paragraph.add_run(part)

if __name__ == '__main__':
    create_docx_from_markdown(
        'LAPORAN_KESELAMATAN_WORDPRESS_MBIKEDAH.md',
        'LAPORAN_KESELAMATAN_WORDPRESS_MBIKEDAH.docx'
    )
    
    # Also convert executive summary
    try:
        create_docx_from_markdown(
            'RINGKASAN_EKSEKUTIF_KESELAMATAN.md',
            'RINGKASAN_EKSEKUTIF_KESELAMATAN.docx'
        )
    except FileNotFoundError:
        print("Ringkasan eksekutif tidak ditemui, hanya laporan utama ditukar.")

